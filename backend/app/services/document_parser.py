"""
Document Parser Service

Handles extraction of text from PDF and DOCX files using:
- PyMuPDF (fitz) for PDF files - fast, structure-preserving
- python-docx for DOCX files
"""

import logging
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, List

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class ParsedSection:
    """A section of parsed text with metadata."""
    text: str
    page_number: Optional[int] = None
    section_type: Optional[str] = None  # heading, paragraph, list, etc.


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    raw_text: str
    sections: List[ParsedSection]
    page_count: int
    word_count: int
    file_type: str


class DocumentParser:
    """
    Handles PDF, DOCX parsing for text extraction.

    Usage:
        parser = DocumentParser()
        result = parser.parse("/path/to/document.pdf")
        print(result.raw_text)
    """

    SUPPORTED_TYPES = {".pdf", ".docx", ".doc"}

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document and extract text.

        Args:
            file_path: Path to the document file

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix in {".docx", ".doc"}:
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """
        Parse PDF using PyMuPDF.

        PyMuPDF is chosen for:
        - Speed (0.12s/doc average)
        - Structure preservation
        - OCR support if needed
        """
        logger.info(f"Parsing PDF: {path}")

        sections: List[ParsedSection] = []
        all_text_parts: List[str] = []

        try:
            doc = fitz.open(str(path))
            page_count = len(doc)

            for page_num, page in enumerate(doc, start=1):
                # Extract text with layout preservation
                text = page.get_text("text")

                if text.strip():
                    sections.append(ParsedSection(
                        text=text,
                        page_number=page_num,
                        section_type="page"
                    ))
                    all_text_parts.append(text)

            doc.close()

            raw_text = "\n\n".join(all_text_parts)
            word_count = len(raw_text.split())

            logger.info(f"PDF parsed: {page_count} pages, {word_count} words")

            return ParsedDocument(
                raw_text=raw_text,
                sections=sections,
                page_count=page_count,
                word_count=word_count,
                file_type="pdf"
            )

        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
            raise

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """
        Parse DOCX using python-docx.
        """
        logger.info(f"Parsing DOCX: {path}")

        sections: List[ParsedSection] = []
        all_text_parts: List[str] = []

        try:
            doc = DocxDocument(str(path))

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Detect section type based on style
                    section_type = "paragraph"
                    if para.style and para.style.name:
                        style_name = para.style.name.lower()
                        if "heading" in style_name:
                            section_type = "heading"
                        elif "title" in style_name:
                            section_type = "title"
                        elif "list" in style_name:
                            section_type = "list"

                    sections.append(ParsedSection(
                        text=text,
                        section_type=section_type
                    ))
                    all_text_parts.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        sections.append(ParsedSection(
                            text=row_text,
                            section_type="table"
                        ))
                        all_text_parts.append(row_text)

            raw_text = "\n\n".join(all_text_parts)
            word_count = len(raw_text.split())

            # DOCX doesn't have pages in the same way as PDF
            page_count = max(1, word_count // 500)  # Rough estimate

            logger.info(f"DOCX parsed: ~{page_count} pages, {word_count} words")

            return ParsedDocument(
                raw_text=raw_text,
                sections=sections,
                page_count=page_count,
                word_count=word_count,
                file_type="docx"
            )

        except Exception as e:
            logger.error(f"Error parsing DOCX {path}: {e}")
            raise

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if a file type is supported."""
        suffix = Path(filename).suffix.lower()
        return suffix in DocumentParser.SUPPORTED_TYPES

    @staticmethod
    def get_file_type(filename: str) -> str:
        """Get the file type from filename."""
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        elif suffix in {".docx", ".doc"}:
            return "docx"
        return "unknown"
